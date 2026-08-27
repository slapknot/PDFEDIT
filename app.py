import streamlit as st
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import tempfile
import os
import io

st.set_page_config(page_title="PDF to Word", page_icon="📄", layout="centered")

st.markdown("""
    <style>
    #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
    .stButton>button { background-color: #10B981; color: white; font-weight: bold; border-radius: 8px; padding: 0.6rem; }
    .stButton>button:hover { background-color: #059669; color: white; }
    </style>
""", unsafe_allow_html=True)

st.title("📄 PDF to Word (Pro Engine)")
st.write("แปลงไฟล์รักษารูปแบบ (Layout) และแทรกรูปภาพตรงตำแหน่งเดิม")

uploaded_file = st.file_uploader("📂 ลากไฟล์ PDF มาวางที่นี่", type=["pdf"])

if uploaded_file is not None:
    if st.button("🚀 เริ่มแปลงไฟล์ (ระบบวิเคราะห์พิกัด)", use_container_width=True):
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(uploaded_file.getvalue())
            pdf_path = tmp_pdf.name
            
        marked_pdf_path = pdf_path.replace(".pdf", "_marked.pdf")
        docx_path = pdf_path.replace(".pdf", ".docx")
        
        try:
            # ========================================================
            # ขั้นตอนที่ 1: PyMuPDF สแกนพิกัดแบบ Deep Scan (X-Ray)
            # ========================================================
            with st.spinner("1/3 กำลังสแกนพิกัด กราฟ และแบบแปลน..."):
                doc = fitz.open(pdf_path)
                img_dict = {}
                img_count = 0
                
                for page in doc:
                    rects_to_capture = []
                    
                    # 1.1 สแกนหา Raster Images (รูปภาพปกติที่ฝังอยู่)
                    for img in page.get_image_info():
                        rects_to_capture.append(fitz.Rect(img["bbox"]))
                        
                    # 1.2 สแกนหา Vector Drawings (ลายเส้น, กราฟ, แบบแปลน)
                    for d in page.get_drawings():
                        r = d["rect"]
                        # กรองเส้นขอบกระดาษ (ใหญ่เกิน) และเส้นคั่นตาราง (เล็กเกิน) ออก
                        if 15 < r.width < (page.rect.width * 0.9) and 15 < r.height < (page.rect.height * 0.9):
                            rects_to_capture.append(r)
                            
                    # 1.3 ยุบรวมพิกัดที่อยู่ติดกัน (Cluster Merging)
                    merged_rects = []
                    for r in rects_to_capture:
                        # ขยายกรอบค้นหา 10 พิกเซลรอบทิศทาง ถ้าชนกันให้ถือว่าเป็นรูปเดียวกัน
                        r_exp = r + (-10, -10, 10, 10) 
                        intersecting = [i for i, mr in enumerate(merged_rects) if r_exp.intersects(mr)]
                        
                        if not intersecting:
                            merged_rects.append(r)
                        else:
                            # ยุบรวมกรอบที่ซ้อนทับกันให้เป็นก้อนเดียว
                            first_idx = intersecting[0]
                            merged_rects[first_idx] = merged_rects[first_idx] | r
                            for i in reversed(intersecting[1:]):
                                merged_rects[first_idx] = merged_rects[first_idx] | merged_rects[i]
                                del merged_rects[i]
                                
                    # 1.4 ถ่ายรูป ถมขาว และฝังรหัสลับ
                    for rect in merged_rects:
                        # ป้องกัน Error รูปขนาด 0
                        if rect.width <= 0 or rect.height <= 0: 
                            continue
                        
                        # ถ่ายรูปพิกัดนั้นแบบชัดๆ (ซูม 2 เท่า)
                        pix = page.get_pixmap(clip=rect, matrix=fitz.Matrix(2.0, 2.0))
                        marker = f"[[IMG_{img_count}]]"
                        
                        img_dict[marker] = {
                            "bytes": pix.tobytes("png"),
                            "width": rect.width
                        }
                        
                        # ถมสีขาวทับรูปเดิมใน PDF และพิมพ์รหัสลับ (สีแดงให้หาเจอชัวร์ๆ)
                        page.draw_rect(rect, color=(1,1,1), fill=(1,1,1))
                        page.insert_text((rect.x0 + 2, rect.y0 + 12), marker, fontsize=10, color=(1,0,0))
                        img_count += 1
                        
                # เซฟ PDF ฉบับที่มีแต่ตัวหนังสือและรหัสลับ
                doc.save(marked_pdf_path)
                doc.close()

            # ========================================================
            # ขั้นตอนที่ 2: ใช้ pdf2docx จัด Format และ Layout
            # ========================================================
            with st.spinner("2/3 กำลังสร้างโครงสร้าง Word และจัด Layout..."):
                cv = Converter(marked_pdf_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()

            # ========================================================
            # ขั้นตอนที่ 3: เปิดไฟล์ Word หาตำแหน่งรหัสลับ แล้ววางรูปลงไป!
            # ========================================================
            with st.spinner("3/3 กำลังประกอบรูปภาพลงในตำแหน่งเดิม..."):
                word_doc = Document(docx_path)
                
                def replace_markers(paragraphs):
                    for p in paragraphs:
                        for marker, img_data in img_dict.items():
                            if marker in p.text:
                                # เคลียร์รหัสลับทิ้ง
                                p.text = "" 
                                # วางรูปภาพลงไปแทนที่
                                run = p.add_run()
                                img_stream = io.BytesIO(img_data["bytes"])
                                # คำนวณขนาดรูปให้พอดี (จำกัดกว้างสุด 6.5 นิ้ว ไม่ให้ล้นหน้า)
                                width_inch = min(img_data["width"] / 72.0, 6.5)
                                run.add_picture(img_stream, width=Inches(width_inch))

                # สแกนหาในข้อความปกติ
                replace_markers(word_doc.paragraphs)
                
                # สแกนหาในตาราง (pdf2docx มักใช้ตารางซ่อนเส้นเพื่อทำเลย์เอาต์)
                for table in word_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_markers(cell.paragraphs)
                            
                word_doc.save(docx_path)

            # ========================================================
            # สำเร็จ! ส่งไฟล์ Word กลับไปให้ผู้ใช้ดาวน์โหลด
            # ========================================================
            with open(docx_path, "rb") as f:
                docx_bytes = f.read()
                
            st.success(f"✅ แปลงไฟล์และแทรกรูปภาพสำเร็จ ({img_count} รูป)")
            
            st.download_button(
                label="📥 ดาวน์โหลดไฟล์ Word (Layout + รูปภาพครบ)",
                data=docx_bytes,
                file_name=uploaded_file.name.rsplit(".", 1)[0] + "_Perfect.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
            
        finally:
            # เคลียร์ไฟล์ขยะ
            for path in [pdf_path, marked_pdf_path, docx_path]:
                if os.path.exists(path):
                    os.remove(path)
